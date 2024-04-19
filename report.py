# -*- coding: utf-8 -*-
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from network import SSH


def reports(commands):
    doc = Document()

    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.styles['Normal'].font.size = Pt(11)

    title = doc.add_heading('Emergency Response Reports', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sshd = SSH()
    sshd.connect()
    for i in commands:
        result = sshd.exec_bash(i)
        doc.add_heading(i, level=1)

        if len(result) > 1:
            doc.add_paragraph(result)
        else:
            doc.add_paragraph('Not Found.')

    filename = f"reports {datetime.now().strftime('%Y-%m-%d %H-%M-%S')}.docx"
    doc.save(filename)
    sshd.close()
